import json
import re
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
BOOK_DIR = ROOT / "book"
CHAPTERS_DIR = BOOK_DIR / "chapters"
DRAFTS_DIR = BOOK_DIR / "drafts" / "chapters"
CONFIG_DIR = BOOK_DIR / "config"
BUILD_DIR = BOOK_DIR / "build"

FIGURE_SLOT_RE = re.compile(r"<!--\s*FIGURE_SLOT:\s*([a-zA-Z0-9._-]+)\s*-->")
VISUAL_RE = re.compile(
    r"<!--\s*VISUAL:\s*([a-zA-Z0-9._-]+)\s*\|\s*id:([a-zA-Z0-9._-]+)\s*\|\s*purpose:(.*?)-->",
    re.IGNORECASE,
)
SECTION_RE = re.compile(r"^(Introduction|Chapter\s+\d+|Epilogue)\b", re.IGNORECASE)
CHAPTER_NUMBER_RE = re.compile(r"^Chapter\s+(\d+)\b", re.IGNORECASE)


def slugify(text):
    slug = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
    return slug or "section"


def chapter_filename(chapter_number, slug):
    return f"{chapter_number:02d}-{slug}.md"


def find_book_sections(docx_path):
    document = Document(str(docx_path))
    paragraphs = [p.text.strip() for p in document.paragraphs]
    markers = []

    for idx, text in enumerate(paragraphs):
        if not text or not SECTION_RE.match(text):
            continue

        section = {"index": idx, "title": text}
        chapter_match = CHAPTER_NUMBER_RE.match(text)

        if text.lower().startswith("introduction"):
            section["kind"] = "introduction"
            section["chapter"] = 0
            section["slug"] = "introduction"
        elif text.lower().startswith("epilogue"):
            section["kind"] = "epilogue"
            section["chapter"] = 13
            section["slug"] = "epilogue"
        elif chapter_match:
            chapter_number = int(chapter_match.group(1))
            title_part = text.split(":", 1)[1].strip() if ":" in text else text
            section["kind"] = "chapter"
            section["chapter"] = chapter_number
            section["slug"] = slugify(title_part)
        else:
            continue

        if (
            markers
            and section["kind"] == "chapter"
            and markers[-1]["kind"] == "chapter"
            and markers[-1]["chapter"] == section["chapter"]
            and section["index"] == markers[-1]["index"] + 1
        ):
            markers[-1] = section
        else:
            markers.append(section)

    return markers


def parse_front_matter(text):
    if not text.startswith("---\n"):
        return {}, text

    parts = text.split("\n---\n", 1)
    if len(parts) != 2:
        return {}, text

    raw_meta, body = parts
    meta_lines = raw_meta.splitlines()[1:]
    meta = {}
    current_key = None

    for line in meta_lines:
        if not line.strip():
            continue
        if line.startswith("  - ") and current_key:
            meta.setdefault(current_key, []).append(line[4:].strip())
            continue
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        current_key = key.strip()
        value = value.strip()
        meta[current_key] = value if value else []

    return meta, body


def read_text(path):
    return Path(path).read_text(encoding="utf-8")


def asset_to_markdown(asset_id, asset, issues):
    source_path = asset.get("source_path", "")
    if not source_path:
        issues.append(f"Asset '{asset_id}' has no source_path.")
        return f"\n[ASSET MISSING PATH: {asset_id}]\n"

    alt = asset.get("alt_text", asset_id)
    caption = asset.get("caption", "")
    footer = asset.get("footer", "")
    reference = asset.get("reference", "")
    blocks = [f"![{alt}]({source_path})"]

    if caption:
        blocks.append(f"*{caption}*")
    if footer:
        blocks.append(footer)
    if reference:
        blocks.append(f"Source: {reference}")

    return "\n" + "\n\n".join(blocks) + "\n"


def replace_figure_slots(body, figures, assets, issues):
    def repl(match):
        slot = match.group(1)
        figure = figures.get(slot)
        if not figure:
            issues.append(f"Unresolved figure slot: {slot}")
            return f"\n[FIGURE MISSING: {slot}]\n"

        return asset_to_markdown(
            slot,
            {
                "source_path": figure.get("path", ""),
                "alt_text": figure.get("alt", slot),
                "caption": figure.get("caption", ""),
                "footer": figure.get("footer", ""),
                "reference": figure.get("reference", ""),
            },
            issues,
        )

    def visual_repl(match):
        visual_kind = match.group(1)
        asset_id = match.group(2)
        _purpose = match.group(3).strip()
        asset = assets.get(asset_id)
        if not asset:
            issues.append(f"Unresolved visual asset id: {asset_id}")
            return f"\n[VISUAL MISSING: {visual_kind} | {asset_id}]\n"
        return asset_to_markdown(asset_id, asset, issues)

    body = FIGURE_SLOT_RE.sub(repl, body)
    body = VISUAL_RE.sub(visual_repl, body)
    return body


def validate_chapter_file(chapter_path, figures, assets):
    text = read_text(chapter_path)
    meta, body = parse_front_matter(text)
    issues = []
    for field in ("chapter", "slug", "title", "status"):
        if field not in meta:
            issues.append(f"{chapter_path.name}: missing metadata '{field}'")

    for slot in FIGURE_SLOT_RE.findall(body):
        if slot not in figures:
            issues.append(f"{chapter_path.name}: unresolved figure slot '{slot}'")

    for _visual_kind, asset_id, _purpose in VISUAL_RE.findall(body):
        if asset_id not in assets:
            issues.append(f"{chapter_path.name}: unresolved visual asset '{asset_id}'")

    return issues


def render_markdown_book(chapter_paths, figures, assets):
    rendered = []
    for chapter_path in chapter_paths:
        text = read_text(chapter_path)
        _, body = parse_front_matter(text)
        issues = []
        rendered_body = replace_figure_slots(body.strip(), figures, assets, issues)
        rendered.append(rendered_body.strip())
    return "\n\n---\n\n".join(part for part in rendered if part)


def add_markdown_to_docx(doc, markdown_text):
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("!["):
            match = re.match(r"!\[[^\]]*\]\(([^)]+)\)", line)
            if match:
                image_path = Path(match.group(1))
                resolved = image_path if image_path.is_absolute() else ROOT / image_path
                if resolved.exists():
                    try:
                        doc.add_picture(str(resolved))
                    except Exception:
                        doc.add_paragraph(f"[IMAGE NOT EMBEDDED: {resolved.name}]")
                else:
                    doc.add_paragraph(f"[IMAGE MISSING: {match.group(1)}]")
        elif line.startswith("*") and line.endswith("*") and len(line) > 2:
            doc.add_paragraph(line.strip("*")).italic = True
        else:
            doc.add_paragraph(raw_line)


def build_docx(chapter_paths, figures, assets, output_path):
    markdown = render_markdown_book(chapter_paths, figures, assets)
    document = Document()
    add_markdown_to_docx(document, markdown)
    document.save(str(output_path))


def load_figures(figures_path):
    if not figures_path.exists():
        return {}
    return json.loads(figures_path.read_text(encoding="utf-8"))


def load_assets(assets_path):
    if not assets_path.exists():
        return {}
    return json.loads(assets_path.read_text(encoding="utf-8"))


def load_book_config(config_path):
    if not config_path.exists():
        return {"chapters": []}
    return json.loads(config_path.read_text(encoding="utf-8"))


def write_text(path, text):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def extract_sections_from_docx(docx_path):
    document = Document(str(docx_path))
    paragraphs = [p.text.strip() for p in document.paragraphs]
    markers = find_book_sections(docx_path)
    extracted = []

    for i, marker in enumerate(markers):
        start = marker["index"]
        end = markers[i + 1]["index"] if i + 1 < len(markers) else len(paragraphs)
        content = [p for p in paragraphs[start:end] if p]
        extracted.append({"meta": marker, "content": content})

    return extracted


def write_chapter_files(sections):
    config = {"chapters": []}

    for section in sections:
        meta = section["meta"]
        content = section["content"]
        chapter_num = meta["chapter"]
        slug = meta["slug"]
        title = meta["title"]
        status = "draft"

        filename = chapter_filename(chapter_num, slug)
        path = CHAPTERS_DIR / filename

        lines = [
            "---",
            f"chapter: {chapter_num:02d}",
            f"slug: {slug}",
            f"title: {title}",
            f"status: {status}",
            "source: Lean and Agile.docx",
            "---",
            "",
        ]

        for idx, paragraph in enumerate(content):
            if idx == 0:
                lines.append(f"# {paragraph}")
                continue
            lines.append(paragraph)
            lines.append("")

        if chapter_num == 1:
            lines.append("<!-- FIGURE_SLOT: hidden-costs-of-rework -->")
            lines.append("")
            lines.append(
                "<!-- EDITORIAL_NOTE: Connect the hidden factory concept to a measurable rework-rate formula. -->"
            )
        elif chapter_num == 5:
            lines.append(
                "<!-- EDITORIAL_NOTE: Add a measurement section covering velocity, cycle time percentiles, commitment accuracy, and lead time. -->"
            )
        elif chapter_num == 6:
            lines.append(
                "<!-- EDITORIAL_NOTE: Expand with a worked example and ITIL metrics such as MTTR, change success rate, and availability. -->"
            )

        write_text(path, "\n".join(lines).strip() + "\n")
        config["chapters"].append(
            {
                "chapter": chapter_num,
                "slug": slug,
                "title": title,
                "path": str(path.relative_to(ROOT)).replace("\\", "/"),
            }
        )

    write_text(CONFIG_DIR / "book.json", json.dumps(config, indent=2))

    figures = {
        "hidden-costs-of-rework": {
            "path": "book/assets/images/hidden-costs-of-rework.png",
            "alt": "Hidden costs of rework",
            "caption": "A visual reminder that rework creates invisible cost throughout the system.",
            "status": "ready",
        }
    }
    write_text(CONFIG_DIR / "figures.json", json.dumps(figures, indent=2))


def chapter_paths_from_config(book_config):
    paths = []
    for item in sorted(book_config["chapters"], key=lambda row: row["chapter"]):
        paths.append(ROOT / item["path"])
    return paths


def draft_chapter_paths_from_config(book_config):
    paths = []
    for item in sorted(book_config["chapters"], key=lambda row: row["chapter"]):
        draft_path = DRAFTS_DIR / Path(item["path"]).name
        if draft_path.exists():
            paths.append(draft_path)
        else:
            paths.append(ROOT / item["path"])
    return paths


def build_outputs():
    book_config = load_book_config(CONFIG_DIR / "book.json")
    figures = load_figures(CONFIG_DIR / "figures.json")
    assets = load_assets(CONFIG_DIR / "assets.json")
    chapter_paths = chapter_paths_from_config(book_config)
    draft_paths = draft_chapter_paths_from_config(book_config)

    issues = []
    for chapter_path in chapter_paths:
        issues.extend(validate_chapter_file(chapter_path, figures, assets))
    for chapter_path in draft_paths:
        issues.extend(validate_chapter_file(chapter_path, figures, assets))

    combined_markdown = render_markdown_book(chapter_paths, figures, assets)
    draft_markdown = render_markdown_book(draft_paths, figures, assets)
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    write_text(BUILD_DIR / "Lean-and-Agile.md", combined_markdown)
    write_text(BUILD_DIR / "Lean-and-Agile-draft.md", draft_markdown)
    build_docx(chapter_paths, figures, assets, BUILD_DIR / "Lean-and-Agile.docx")
    build_docx(draft_paths, figures, assets, BUILD_DIR / "Lean-and-Agile-draft.docx")
    write_text(BUILD_DIR / "validation-report.txt", "\n".join(issues) if issues else "No validation issues detected.\n")


def main():
    if len(sys.argv) < 2:
        print("Usage: build_book.py [extract|build|validate]")
        return 1

    command = sys.argv[1].lower()
    if command == "extract":
        sections = extract_sections_from_docx(ROOT / "Lean and Agile.docx")
        write_chapter_files(sections)
        return 0
    if command == "build":
        build_outputs()
        return 0
    if command == "validate":
        book_config = load_book_config(CONFIG_DIR / "book.json")
        figures = load_figures(CONFIG_DIR / "figures.json")
        assets = load_assets(CONFIG_DIR / "assets.json")
        chapter_paths = chapter_paths_from_config(book_config)
        issues = []
        for chapter_path in chapter_paths:
            issues.extend(validate_chapter_file(chapter_path, figures, assets))
        print("\n".join(issues) if issues else "No validation issues detected.")
        return 0

    print(f"Unknown command: {command}")
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
